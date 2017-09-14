"""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""

    Author: Rajagopal Senthil Kumar
    Created Date:  07-Sep-2017
    Modified Date: 07-Sep-2017
    Purpose: Python Program to read document and extract the texts/images out

"""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""""


import docx
import sys
from docx.shape import InlineShape, InlineShapes
from docx.enum.shape import WD_INLINE_SHAPE
from docx import Document
from internal import Internal


class DocReader(Internal):

    def __init__(self,file_name="",extract_yes=0):
        self.file_name=file_name
        self.inline_shapes_count={'no_of_pictures':0,'no_of_linkedpictures':0,'no_of_charts':0,'no_of_smartarts':0,'no_of_notimplemented':0}
        if file_name !="": 
            self.word_file=open(self.file_name,'rb+')
            self.word_document=Document(self.word_file)
            if extract_yes==1:
                self.countInlineShapes()

    def countInlineShapes(self):
        word_shapes = self.word_document.inline_shapes
        word_shapes_count=0

        for word_shape in word_shapes:
            word_shapes_count+=1
            if word_shape.type==WD_INLINE_SHAPE.PICTURE:
               self.inline_shapes_count['no_of_pictures']+=1
                  
            elif word_shape.type==WD_INLINE_SHAPE.LINKED_PICTURE:
                self.inline_shapes_count['no_of_linkedpictures']+=1
            elif word_shape.type==WD_INLINE_SHAPE.CHART:
                self.inline_shapes_count['no_of_charts'] +=1
            elif word_shape.type==WD_INLINE_SHAPE.SMART_ART:
                self.inline_shapes_count['no_of_smartarts'] +=1
            elif word_shape.type==WD_INLINE_SHAPE.NOT_IMPLEMENTED:
                self.inline_shapes_count['no_of_notimplemented'] +=1

 

#End of this Program. Other counting and extract function will be added in future######################################






        




    
#print (word_shapes_count)

#def getText(filename):
#    doc = docx.Document(filename)
#    fullText = []
#    for para in doc.paragraphs:
#        fullText.append(para.text)
#    return '\n'.join(fullText)




#file01= "C:\\Program360\\Sample\\MySample01.docx"

#text1 =getText(file01)
#doc1 = Document(file01)

#tables = doc1.tables

#for table in tables:
#    for row in table.rows:
#        for cell in row.cells:
#            for paragraph in cell.paragraphs:
#                print(paragraph.text)

#print(text1)


#inline_shapes = doc1.body.InlineShapes
#for inline_shape in inline_shapes:
#    print(inline_shape.type)









